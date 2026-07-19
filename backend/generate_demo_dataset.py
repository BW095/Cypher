import os
import random
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from faker import Faker

fake = Faker()
Faker.seed(42)
random.seed(42)

BASE_DIR = "/home/bw/CODES/demo_dataset_large"

folders = ["manuals", "logs", "maintenance", "incidents"]
for folder in folders:
    os.makedirs(os.path.join(BASE_DIR, folder), exist_ok=True)

# Equipments
turbines = [f"GT-{i}" for i in range(201, 206)]
pumps = [f"PUMP-{i}" for i in range(101, 111)]
compressors = [f"COMP-{i}" for i in range(301, 306)]
equipments = turbines + pumps + compressors

# 1. Manuals (10 long manuals)
def generate_manual(eq_name, filename):
    doc = Document()
    doc.add_heading(f'OEM Manual - {eq_name}', 0)
    for i in range(1, 11): # 10 sections
        doc.add_heading(f'{i}. {fake.catch_phrase()}', level=1)
        for _ in range(5): # 5 paragraphs per section
            doc.add_paragraph(fake.text(max_nb_chars=1000))
        if i == 5:
            doc.add_heading('Critical Operating Limits', level=2)
            doc.add_paragraph(f'Vibration Trip Limit for {eq_name}: {round(random.uniform(8.0, 15.0), 1)} mm/s')
            doc.add_paragraph(f'Temperature Warning Limit: {random.randint(85, 95)}°C')
    doc.save(os.path.join(BASE_DIR, 'manuals', filename))

for eq in turbines + compressors:
    generate_manual(eq, f"{eq}_OEM_Manual_Rev3.docx")

# 2. Logs - Shift Logs (30 files)
start_date = datetime(2026, 1, 1)
for day in range(30):
    current_date = start_date + timedelta(days=day)
    date_str = current_date.strftime('%Y-%m-%d')
    with open(os.path.join(BASE_DIR, 'logs', f'shift_log_{date_str}.txt'), 'w') as f:
        f.write(f"SHIFT LOG - {date_str}\n")
        f.write("="*40 + "\n\n")
        for shift in ['Morning (06:00-14:00)', 'Afternoon (14:00-22:00)', 'Night (22:00-06:00)']:
            f.write(f"--- Shift: {shift} ---\n")
            f.write(f"Operator: {fake.name()}\n")
            f.write("Notes:\n")
            for _ in range(3):
                f.write(f"- {fake.sentence()}\n")
            # Inject a realistic note occasionally
            if random.random() < 0.2:
                target_eq = random.choice(equipments)
                f.write(f"- Field operator noted unusual vibration on {target_eq}. Monitoring closely.\n")
            f.write("\n")

# 2. Logs - Sensor Data (5 large files)
for eq in random.sample(equipments, 5):
    # Generate 43,200 rows (30 days, minute-by-minute)
    dates = pd.date_range(start="2026-01-01", periods=43200, freq="min")
    temp_base = random.randint(50, 70)
    vib_base = random.uniform(1.0, 3.0)
    
    # Introduce some noise and trend
    temps = [temp_base + (i/10000.0) + random.uniform(-2, 2) for i in range(43200)]
    vibs = [vib_base + (i/20000.0) + random.uniform(-0.5, 0.5) for i in range(43200)]
    
    df = pd.DataFrame({
        "Timestamp": dates,
        "Equipment_ID": [eq]*43200,
        "Bearing_Temp_C": temps,
        "Vibration_RMS_mms": vibs,
        "Line_Pressure_PSI": [random.uniform(100, 110) for _ in range(43200)]
    })
    df.to_excel(os.path.join(BASE_DIR, 'logs', f'sensor_data_{eq}.xlsx'), index=False)

# 3. Maintenance - Work Orders (500 rows)
wo_numbers = [f"WO-{10000 + i}" for i in range(500)]
wo_dates = [start_date + timedelta(days=random.randint(0, 30)) for _ in range(500)]
wo_eq = [random.choice(equipments) for _ in range(500)]
wo_desc = [fake.sentence() for _ in range(500)]
wo_status = [random.choice(["Open", "Closed", "In Progress", "Deferred"]) for _ in range(500)]
wo_assigned = [fake.name() for _ in range(500)]

# Inject a specific storyline for the demo
wo_numbers[42] = "WO-2041"
wo_dates[42] = datetime(2026, 1, 10)
wo_eq[42] = "GT-201"
wo_desc[42] = "Emergency replacement of compressor bearing due to high vibration trip."
wo_status[42] = "Closed"
wo_assigned[42] = "Technician Alex"

df_wo = pd.DataFrame({
    "WO_Number": wo_numbers,
    "Date": wo_dates,
    "Equipment": wo_eq,
    "Description": wo_desc,
    "Status": wo_status,
    "Assigned_To": wo_assigned
})
df_wo.to_excel(os.path.join(BASE_DIR, 'maintenance', 'master_work_orders.xlsx'), index=False)

# 3. Maintenance Reports (15 docx files)
def generate_maintenance_report(wo_num, eq, filename):
    doc = Document()
    doc.add_heading(f'Maintenance Report: {wo_num}', 0)
    doc.add_paragraph(f'Equipment: {eq}')
    doc.add_paragraph(f'Technician: {fake.name()}')
    doc.add_heading('Task Description', level=1)
    doc.add_paragraph(fake.paragraph(nb_sentences=5))
    doc.add_heading('Action Taken', level=1)
    for _ in range(5):
        doc.add_paragraph(f"- {fake.sentence()}")
    doc.add_heading('Post-Maintenance Testing', level=1)
    doc.add_paragraph(fake.paragraph(nb_sentences=4))
    doc.save(os.path.join(BASE_DIR, 'maintenance', filename))

for i in range(15):
    target = random.choice(equipments)
    generate_maintenance_report(f"WO-{10500+i}", target, f"maintenance_report_{target}_{i}.docx")

# Add the specific story maintenance report
generate_maintenance_report("WO-2041", "GT-201", "maintenance_report_GT201_Emergency.docx")

# 4. Incident Reports (10 docx files)
def generate_incident_report(eq, filename, is_story=False):
    doc = Document()
    doc.add_heading(f'Incident / RCA Report - {eq}', 0)
    doc.add_paragraph(f'Investigator: {fake.name()}')
    doc.add_heading('Incident Description', level=1)
    if is_story:
        doc.add_paragraph('GT-201 tripped on Jan 09 due to compressor bearing failure, resulting in 48 hours of downtime. ' + fake.paragraph(nb_sentences=10))
    else:
        doc.add_paragraph(fake.paragraph(nb_sentences=10))
    
    doc.add_heading('Root Cause Analysis (5 Whys)', level=1)
    for i in range(1, 6):
        doc.add_paragraph(f"Why {i}: {fake.sentence()}")
        
    doc.add_heading('Corrective Actions', level=1)
    for _ in range(3):
        doc.add_paragraph(f"- {fake.sentence()}")
    doc.save(os.path.join(BASE_DIR, 'incidents', filename))

for i in range(9):
    generate_incident_report(random.choice(equipments), f"incident_report_{i}.docx")

# Specific story RCA
generate_incident_report("GT-201", "root_cause_analysis_GT201.docx", is_story=True)

print("Massive dataset generated successfully at", BASE_DIR)
