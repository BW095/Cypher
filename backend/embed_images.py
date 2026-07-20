import os
import glob
import shutil
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

# Artifact dir
ARTIFACT_DIR = "/home/bw/.gemini/antigravity-ide/brain/233866c7-00f2-4dbc-9d54-782fbb75f830"
DEMO_DIR = "/home/bw/CODES/demo"

# Create images folder
img_dir = os.path.join(DEMO_DIR, "images")
os.makedirs(img_dir, exist_ok=True)

# Copy and rename images
image_map = {
    "gas_turbine_diagram": None,
    "damaged_bearing": None,
    "control_room_alarm": None,
    "lube_oil_filter": None
}

for img_file in glob.glob(os.path.join(ARTIFACT_DIR, "*.png")):
    filename = os.path.basename(img_file)
    for key in image_map.keys():
        if filename.startswith(key):
            dest = os.path.join(img_dir, f"{key}.png")
            shutil.copy(img_file, dest)
            image_map[key] = dest

# 1. Embed in GT-201_OEM_Manual_Rev3.docx
manual_path = os.path.join(DEMO_DIR, "manuals", "GT-201_OEM_Manual_Rev3.docx")
doc = Document(manual_path)
doc.add_heading('Appendix A: Technical Diagrams', level=1)
doc.add_paragraph('Below is the cross-sectional engineering diagram for the GT-201 unit.')
doc.add_picture(image_map["gas_turbine_diagram"], width=Inches(6.0))
doc.save(manual_path)

# 2. Embed in root_cause_analysis_GT201.docx
rca_path = os.path.join(DEMO_DIR, "incidents", "root_cause_analysis_GT201.docx")
doc2 = Document(rca_path)
doc2.add_heading('Photographic Evidence', level=1)
doc2.add_paragraph('Exhibit A: Damaged Compressor Bearing showing severe scoring.')
doc2.add_picture(image_map["damaged_bearing"], width=Inches(5.0))
doc2.add_paragraph('Exhibit B: Crushed and clogged FS-20 Lube Oil Filter.')
doc2.add_picture(image_map["lube_oil_filter"], width=Inches(5.0))
doc2.save(rca_path)

# 3. Create GT-201_Emergency_Shutdown_Log.pdf
pdf_path = os.path.join(DEMO_DIR, "incidents", "GT-201_Emergency_Shutdown_Log.pdf")
class PDF(FPDF):
    def header(self):
        self.set_font('helvetica', 'B', 15)
        self.cell(80)
        self.cell(30, 10, 'Emergency Shutdown Log - GT-201', 0, 0, 'C')
        self.ln(20)
    
pdf = PDF()
pdf.add_page()
pdf.set_font('helvetica', '', 12)
pdf.multi_cell(0, 10, 'Date: Jan 09, 2026\nTime: 18:42 PM\nTrigger: High Vibration Alarm on Compressor Bearing\n\nBelow is the photograph captured by the shift supervisor of the DCS control screen at the exact moment of the trip, showing vibration levels exceeding 12.0 mm/s.')
pdf.ln(10)
pdf.image(image_map["control_room_alarm"], x=10, w=190)
pdf.output(pdf_path)

print("Images successfully embedded and PDF created!")
